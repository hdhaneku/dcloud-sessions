import requests
from docx import Document
import json

# Base URL with placeholder
BASE_URL = "https://dcloud2-rtp.cisco.com/api/sessions/{session_id}?expand=all"

# Your bearer token
# NOTE: This token might expire. You'll need to update it if you encounter authentication errors.
BEARER_TOKEN = "eyJraWQiOiJIUXh2eW4zM3VOSEVQMFd1RUZiNFpMTW9nZzFnMDhTRnJQM2haQ0g3cjJRIiwiYWxnIjoiUlMyNTYifQ.eyJ2ZXIiOjEsImp0aSI6IkFULjZNOE1ueG96MUlPV091MjBLUUVnZHhVdnNsRjhLb3hsOW9SMEJraEk2U2cub2FyM3AzejRjbkVYVEZ5M2k1ZDciLCJpc3MiOiJodHRwczovL2lkLmNpc2NvLmNvbS9vYXV0aDIvZGVmYXVsdCIsImF1ZCI6ImFwaTovL2RlZmF1bHQiLCJpYXQiOjE3NTU4OTM4NjYsImV4cCI6MTc1NTg5NzQ2NiwiY2lkIjoiMmZjMzI2YjItZjMxYy00MTVmLWI3ZTEtYmI3NmU5MDhhNzllIiwidWlkIjoiMDB1ZG9qZ2E5Q09Pc0pEVnM1ZDYiLCJzY3AiOlsib3BlbmlkIiwib2ZmbGluZV9hY2Nlc3MiLCJjY2lfY29pbWVtYmVyT2YiLCJncm91cHMiLCJlbWFpbCIsInByb2ZpbGUiLCJjY2lfYWRtZW1iZXJPZiJdLCJhdXRoX3RpbWUiOjE3NTU3OTE2MTUsImFjY2Vzc19sZXZlbCI6NCwic3ViIjoiaGRoYW5la3VAY2lzY28uY29tIiwiZW1haWxfYWRkcmVzcyI6ImhkaGFuZWt1QGNpc2NvLmNvbSIsImF6cCI6IjJmYzMyNmIyLWYzMWMtNDE1Zi1iN2UxLWJiNzZlOTA4YTc5ZSIsImdyb3VwcyI6WyJFdmVyeW9uZSIsIkZlZGVyYXRlZCBVc2VyIiwiU21hcnRzaGVldC1BY3RpdmVVc2Vycy1TYWxlc09SRyIsInByb3ZpZGVyLWNvbm5lY3Rpdml0eS1UTUUiLCJFbXBsb3llZXMiLCJJbkJpel9BY2Nlc3MiLCJHSEVDX2Npc2NvLW1hZ25ldGljX2xvZ2luIl0sImNjb2lkIjoiaGRoYW5la3UifQ.TFlZ0uR1yeU3iKz-oxE5THpNaXAHVjqx0YboJubZtU8SemD44eRnQtdH3AKlUWC79ok4fDG7CT8dO9_HGcS-bPDmdOJR1L-cWMUp210rAol3FMjoxvs4fO0aJka1oI3P1WS5TZu_dVdKMNPZ59xo09U1UjiuH6pXAwgDrwzX_c2dksuff6gM-bjJuZgd_vEJlWWcllSjAmMWI2p7SseiAD_PKDs-XVYkgWQZx5jK0ag1ooYgt4-Lo7BSzqifpRJT1FEQJW9IkXU64xxYA0U2ePxxXi1uUQQ9NVhD2qLuIetHWApvCIaB7kTXg5tyn7BB5WbIvYbcqj4HZFwSH9OKAQ"

# Headers with authentication
HEADERS = {
    "Authorization": f"Bearer {BEARER_TOKEN}",
    "Accept": "application/json"
}

# Create Word document
doc = Document()

for session_id in range(1250385, 1250434):
    url = BASE_URL.format(session_id=session_id)
    try:
        response = requests.get(url, headers=HEADERS, timeout=10)
        response.raise_for_status() # Raise an exception for HTTP errors (4xx or 5xx)
        data = response.json()

        # Navigate to the 'network' section within 'expand' to get VPN details
        network_info = data.get("expand", {}).get("network", {})

        # Extract VPN details using .get() with a default value
        vpn_user = network_info.get("vpnUserIds", "N/A")
        vpn_pass = network_info.get("vpnPassword", "N/A")
        vpn_server = network_info.get("vpnServer", "N/A")

        # Add to Word doc
        doc.add_heading(f"Session ID: {session_id}", level=1)
        doc.add_paragraph(f"VPN Username: {vpn_user}")
        doc.add_paragraph(f"VPN Password: {vpn_pass}")
        doc.add_paragraph(f"VPN Server: {vpn_server}")

    except requests.exceptions.RequestException as e:
        # Catch requests-specific exceptions for network errors, timeouts, HTTP errors
        doc.add_heading(f"Session ID: {session_id}", level=1)
        doc.add_paragraph(f"Error retrieving session {session_id}: {e}")
    except json.JSONDecodeError as e:
        # Catch JSON decoding errors
        doc.add_heading(f"Session ID: {session_id}", level=1)
        doc.add_paragraph(f"Error decoding JSON for session {session_id}: {e}")
    except Exception as e:
        # Catch any other unexpected errors
        doc.add_heading(f"Session ID: {session_id}", level=1)
        doc.add_paragraph(f"An unexpected error occurred for session {session_id}: {e}")

    doc.add_page_break()

# Save document
output_file = "vpn_sessions_rtp.docx"
doc.save(output_file)

print(f"Word document saved as {output_file}")
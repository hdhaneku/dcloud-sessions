import requests
from docx import Document
import json

def generate_vpn_report(bearer_token: str, start_session_id: int, end_session_id: int, output_filename: str):
    """
    Fetches VPN details for a range of dcloud sessions and saves them to a Word document.

    Args:
        bearer_token (str): The authentication bearer token for the API.
        start_session_id (int): The first session ID in the range to query.
        end_session_id (int): The last session ID in the range to query (inclusive).
        output_filename (str): The name of the Word document to save the report to.
    """
    # Base URL with placeholder
    BASE_URL = "https://dcloud2-rtp.cisco.com/api/sessions/{session_id}?expand=all"

    # Headers with authentication
    HEADERS = {
        "Authorization": f"Bearer {bearer_token}",
        "Accept": "application/json"
    }

    # Create Word document
    doc = Document()

    # Loop through the specified range of session IDs
    for session_id in range(start_session_id, end_session_id + 1):
        url = BASE_URL.format(session_id=session_id)
        try:
            response = requests.get(url, headers=HEADERS, timeout=10)
            response.raise_for_status() # Raise an exception for HTTP errors (4xx or 5xx)
            data = response.json()

            # Navigate to the 'network' section within 'expand' to get VPN details
            network_info = data.get("expand", {}).get("network", {})

            # Extract VPN details using .get() with a default value
            vpn_user = network_info.get("vpnUserIds", "N/A")
            vpn_pass = network_info.get("vpnPassword", "N/A")
            vpn_server = network_info.get("vpnServer", "N/A")

            # Add to Word doc
            doc.add_heading(f"Session ID: {session_id}", level=1)
            doc.add_paragraph(f"VPN Username: {vpn_user}")
            doc.add_paragraph(f"VPN Password: {vpn_pass}")
            doc.add_paragraph(f"VPN Server: {vpn_server}")

        except requests.exceptions.RequestException as e:
            # Catch requests-specific exceptions for network errors, timeouts, HTTP errors
            doc.add_heading(f"Session ID: {session_id}", level=1)
            doc.add_paragraph(f"Error retrieving session {session_id}: {e}")
        except json.JSONDecodeError as e:
            # Catch JSON decoding errors
            doc.add_heading(f"Session ID: {session_id}", level=1)
            doc.add_paragraph(f"Error decoding JSON for session {session_id}: {e}")
        except Exception as e:
            # Catch any other unexpected errors
            doc.add_heading(f"Session ID: {session_id}", level=1)
            doc.add_paragraph(f"An unexpected error occurred for session {session_id}: {e}")

        # Add a page break after each session's report
        doc.add_page_break()

    # Save document
    doc.save(output_filename)
    print(f"Word document saved as {output_filename}")

if __name__ == "__main__":
    print("--- VPN Session Report Configuration ---")
    try:
        # Get user input for configuration variables
        bearer_token = input("Enter your Bearer Token: ")
        start_session_id = int(input("Enter the starting session ID: "))
        end_session_id = int(input("Enter the ending session ID: "))
        output_file = input("Enter the output filename (e.g., vpn_report.docx): ")

        # Call the function to generate the report
        generate_vpn_report(bearer_token, start_session_id, end_session_id, output_file)
    except ValueError:
        print("\nInvalid input. Please enter valid integer values for the session IDs.")
    except Exception as e:
        print(f"\nAn error occurred: {e}")
